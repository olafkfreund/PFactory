"""Tests for PDF / DOCX plan ingestion (issue #2).

These exercise the binary document loaders that flatten an uploaded plan to
text and reuse the existing ``spec_sources`` parsers. They are skipped where
``pypdf`` / ``python-docx`` aren't installed (e.g. the minimal backend venv).
"""

from __future__ import annotations

import io
import sys
from pathlib import Path

import pytest

_BACKEND = Path(__file__).parent.parent / "apps" / "backend"
if str(_BACKEND) not in sys.path:
    sys.path.insert(0, str(_BACKEND))

# Optional deps — skip the whole module if either is missing.
pytest.importorskip("pypdf")
pytest.importorskip("docx")

from plan.ingest.document_loaders import (  # noqa: E402
    DocumentLoadError,
    extract_docx_text,
    extract_pdf_text,
    extract_text,
    ingest_document,
    load_document_text,
)
from spec_sources import SpecFormat  # noqa: E402

# ── fixtures / builders ────────────────────────────────────────────────

def _make_docx(*, title: str, description: str, criteria: list[str]) -> bytes:
    """Build a real .docx with a title heading, body, and a bulleted AC list."""
    import docx

    d = docx.Document()
    d.add_heading(title, level=1)
    if description:
        d.add_paragraph(description)
    d.add_heading("Acceptance Criteria", level=2)
    for c in criteria:
        d.add_paragraph(c, style="List Bullet")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _make_pdf(lines: list[str]) -> bytes:
    """Build a minimal single-page PDF with one text line per entry."""
    text_ops = (
        "BT /F1 12 Tf 72 740 Td 16 TL "
        + " ".join(f"({ln}) Tj T*" for ln in lines)
        + " ET"
    )
    objs = [
        b"<</Type/Catalog/Pages 2 0 R>>",
        b"<</Type/Pages/Kids[3 0 R]/Count 1>>",
        b"<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]"
        b"/Resources<</Font<</F1 4 0 R>>>>/Contents 5 0 R>>",
        b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>",
    ]
    stream = text_ops.encode()
    objs.append(b"<</Length %d>>\nstream\n%s\nendstream" % (len(stream), stream))
    out = b"%PDF-1.4\n"
    offsets: list[int] = []
    for i, o in enumerate(objs, start=1):
        offsets.append(len(out))
        out += b"%d 0 obj\n%s\nendobj\n" % (i, o)
    xref = len(out)
    out += b"xref\n0 %d\n0000000000 65535 f \n" % (len(objs) + 1)
    for off in offsets:
        out += b"%010d 00000 n \n" % off
    out += b"trailer\n<</Size %d/Root 1 0 R>>\nstartxref\n%d\n%%%%EOF" % (
        len(objs) + 1,
        xref,
    )
    return out


# ── extractors ─────────────────────────────────────────────────────────

def test_extract_docx_reconstructs_markdown_structure():
    data = _make_docx(
        title="Payments Service",
        description="Add a refund endpoint.",
        criteria=["User can request a refund", "Refunds are audited"],
    )
    text = extract_docx_text(data)
    assert "# Payments Service" in text
    assert "## Acceptance Criteria" in text
    assert "- User can request a refund" in text
    assert "- Refunds are audited" in text


def test_extract_pdf_text_returns_lines():
    pdf = _make_pdf(["My Plan", "AC#1: User can log in", "AC#2: Session expires"])
    text = extract_pdf_text(pdf)
    assert "AC#1: User can log in" in text
    assert "AC#2: Session expires" in text


# ── end-to-end ingestion ───────────────────────────────────────────────

def test_ingest_document_docx(tmp_path):
    path = tmp_path / "plan.docx"
    path.write_bytes(
        _make_docx(
            title="Search Feature",
            description="Add full-text search.",
            criteria=["Results return in under 200ms", "Typos are tolerated"],
        )
    )
    spec = ingest_document(path)
    assert spec.title == "Search Feature"
    assert spec.source_format is SpecFormat.MARKDOWN
    texts = [c.text for c in spec.criteria]
    assert "Results return in under 200ms" in texts
    assert "Typos are tolerated" in texts


def test_ingest_document_pdf_uses_inline_ac_markers(tmp_path):
    path = tmp_path / "plan.pdf"
    path.write_bytes(
        _make_pdf(
            ["Onboarding Plan", "AC#1: User verifies email", "AC#2: Profile is created"]
        )
    )
    spec = ingest_document(path)
    assert len(spec.criteria) == 2
    assert spec.criteria[0].text == "User verifies email"
    assert spec.criteria[1].text == "Profile is created"


def test_load_document_text_passes_through_markdown(tmp_path):
    md = "# Title\n\n## Acceptance Criteria\n- one\n- two\n"
    path = tmp_path / "plan.md"
    path.write_text(md)
    assert load_document_text(path) == md


# ── error paths ────────────────────────────────────────────────────────

def test_unsupported_suffix_raises():
    with pytest.raises(DocumentLoadError, match="unsupported document type"):
        extract_text(b"\x00\x01", filename="sheet.xlsx")


def test_empty_document_raises(tmp_path):
    import docx

    buf = io.BytesIO()
    docx.Document().save(buf)  # a blank document — no paragraphs at all
    path = tmp_path / "empty.docx"
    path.write_bytes(buf.getvalue())
    with pytest.raises(DocumentLoadError, match="no extractable text"):
        ingest_document(path)


def test_missing_file_raises():
    with pytest.raises(DocumentLoadError, match="could not read"):
        load_document_text("/nonexistent/plan.pdf")
