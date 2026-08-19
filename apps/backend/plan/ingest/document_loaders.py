"""PDF / DOCX document loaders for plan ingestion (issue #2).

PFactory accepts uploaded project plans as ``docx``, ``pdf``, or ``markdown``.
The existing :mod:`spec_sources` module already normalises *text* sources
(markdown / Gherkin / EARS) into a :class:`~spec_sources.NormalizedSpec`. This
module adds the missing binary loaders: it flattens a PDF or DOCX to plain text
and hands that text to the same parsers, so one upload path serves every format.

Design notes:
  * **Lazy imports.** ``pypdf`` / ``python-docx`` are imported inside the
    extractor functions so the rest of the backend (and text-only ingestion)
    never pays for them and they stay optional at import time.
  * **DOCX → markdown.** Word paragraph *styles* carry structure that flat text
    loses. We map ``Heading N`` paragraphs to ``#``-level headings and list
    paragraphs to ``-`` bullets, so a Word plan with an "Acceptance Criteria"
    heading and a bulleted list round-trips into the markdown shape
    :func:`spec_sources.parse_markdown` understands.
  * **Bytes-first.** :func:`extract_text` works on raw bytes + a filename, so the
    portal upload route (issue #4) can ingest an in-memory upload without
    touching disk; :func:`load_document_text` is the path-based convenience.
"""

from __future__ import annotations

import io
import re
from pathlib import Path

from spec_sources import NormalizedSpec, SpecFormat, ingest

__all__ = [
    "DocumentLoadError",
    "extract_docx_text",
    "extract_pdf_text",
    "extract_text",
    "ingest_document",
    "load_document_text",
]

# Suffixes whose bytes are already text — decoded as UTF-8 and parsed directly.
TEXT_SUFFIXES: frozenset[str] = frozenset(
    {".md", ".markdown", ".txt", ".text", ".feature", ".rst", ""}
)
# Binary document suffixes this module can flatten to text.
DOCUMENT_SUFFIXES: frozenset[str] = frozenset({".pdf", ".docx"})


class DocumentLoadError(ValueError):
    """Raised when a document can't be read or yields no extractable text.

    Verified safe to return to the client verbatim (Factory#718): every raise
    site below now describes only the caller's own filename/type/path, or a
    static sentence -- none interpolates a caught exception's text (pypdf /
    python-docx raise a wide, unreviewed variety of messages on malformed
    input, which is why ``except Exception as exc`` further down deliberately
    stops at ``type(exc).__name__`` instead of ``exc``). See
    ``client_errors.client_error``.
    """

    @property
    def client_message(self) -> str:
        return str(self)


# ── extractors ─────────────────────────────────────────────────────────


def extract_pdf_text(data: bytes) -> str:
    """Extract text from PDF bytes, one page per block, joined by newlines."""
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - environment guard
        raise DocumentLoadError(
            "PDF ingestion needs the 'pypdf' package (pip install pypdf)."
        ) from exc
    try:
        reader = PdfReader(io.BytesIO(data))
        pages = [(page.extract_text() or "") for page in reader.pages]
    except Exception as exc:  # pypdf raises a variety of read errors
        # Factory#718: pypdf's own message is not reviewed for what it reveals
        # (it can echo internal object state) -- only its class name crosses
        # the boundary; the full exception is still on `from exc`'s chain for
        # whatever logs the eventual HTTPException.
        raise DocumentLoadError(f"could not read PDF ({type(exc).__name__})") from exc
    return "\n".join(pages).strip()


_HEADING_LEVEL = re.compile(r"heading\s+(\d+)", re.IGNORECASE)


def extract_docx_text(data: bytes) -> str:
    """Extract text from DOCX bytes, reconstructing markdown from paragraph styles.

    ``Heading N`` → ``#``×N heading; any ``List*`` style → ``-`` bullet; table
    cells are appended as plain lines. The result feeds
    :func:`spec_sources.parse_markdown`.
    """
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - environment guard
        raise DocumentLoadError(
            "DOCX ingestion needs the 'python-docx' package (pip install python-docx)."
        ) from exc
    try:
        doc = Document(io.BytesIO(data))
    except Exception as exc:  # python-docx raises a variety of read errors
        # Factory#718: same reasoning as extract_pdf_text -- class name only.
        raise DocumentLoadError(f"could not read DOCX ({type(exc).__name__})") from exc

    lines: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower() if para.style else ""
        heading = _HEADING_LEVEL.search(style)
        if style.startswith("title"):
            lines.append(f"# {text}")
        elif heading:
            level = min(int(heading.group(1)), 6)
            lines.append(f"{'#' * level} {text}")
        elif "list" in style:
            lines.append(f"- {text}")
        else:
            lines.append(text)

    # Table cells carry acceptance criteria too — flatten each row as a line.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    return "\n".join(lines).strip()


def extract_text(data: bytes, *, filename: str) -> str:
    """Flatten an uploaded document (by ``filename`` extension) to plain text.

    Bytes-first entry point for upload streams. Raises
    :class:`DocumentLoadError` for unsupported types.
    """
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return extract_pdf_text(data)
    if suffix == ".docx":
        return extract_docx_text(data)
    if suffix in TEXT_SUFFIXES:
        return data.decode("utf-8", errors="replace")
    raise DocumentLoadError(
        f"unsupported document type: {suffix or '(no extension)'} — "
        f"supported: {', '.join(sorted(DOCUMENT_SUFFIXES | TEXT_SUFFIXES - {''}))}."
    )


def load_document_text(path: str | Path) -> str:
    """Read a plan document from disk and return its extracted text."""
    p = Path(path)
    try:
        data = p.read_bytes()
    except OSError as exc:
        # Factory#718: `p` is the caller's own requested path, safe to echo;
        # the OSError's own text is not reviewed, so class name only.
        raise DocumentLoadError(f"could not read {p}: {type(exc).__name__}") from exc
    return extract_text(data, filename=p.name)


# ── unified ingestion ──────────────────────────────────────────────────


def ingest_document(
    source: str | Path,
    *,
    fmt: SpecFormat | None = None,
    title: str | None = None,
) -> NormalizedSpec:
    """Ingest a plan document (PDF / DOCX / text) into a :class:`NormalizedSpec`.

    Flattens the document to text, then reuses :func:`spec_sources.ingest` for
    format detection and acceptance-criteria parsing. ``fmt`` forces a parser;
    otherwise it is auto-detected from the extracted text.
    """
    p = Path(source)
    text = load_document_text(p)
    if not text.strip():
        raise DocumentLoadError(f"no extractable text in {p.name}")
    return ingest(text, fmt=fmt, filename=p.name, title=title)
